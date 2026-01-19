
import sys
import os
import docx
from docx.oxml.ns import qn
from lxml import etree

sys.path.append("/app")

def inspect_table_drawings(file_path):
    print(f"Inspecting Drawings in Tables for: {file_path}")
    doc = docx.Document(file_path)
    
    count = 0
    for t_idx, tbl in enumerate(doc.tables):
        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, p in enumerate(cell.paragraphs):
                    for child in p._element.iterdescendants():
                        if child.tag.endswith('drawing'):
                            count += 1
                            print(f"\n[Table {t_idx+1} R{r_idx}C{c_idx} P{p_idx}] Found w:drawing")
                            
                            # Check what's inside
                            xml_str = etree.tostring(child, pretty_print=True).decode()
                            
                            has_blip = 'blip' in xml_str and 'embed' in xml_str
                            has_dgm = 'dgm' in xml_str # Diagram/SmartArt
                            has_wsp = 'wsp' in xml_str or 'shapetype' in xml_str # Shapes
                            has_chart = 'c:chart' in xml_str
                            
                            print(f"  - Has Blip (Image): {has_blip}")
                            print(f"  - Has Diagram (SmartArt): {has_dgm}")
                            print(f"  - Has Shapes (Vector): {has_wsp}")
                            print(f"  - Has Chart: {has_chart}")
                            
                            # Check Dimensions
                            extent = None
                            for d in child.iterdescendants():
                                if d.tag.endswith('extent'):
                                    extent = (d.get('cx'), d.get('cy'))
                                    break
                            print(f"  - Extent: {extent}")
                            
                            # Check docPr (Name, Descr)
                            docPr = None
                            for d in child.iterdescendants():
                                if d.tag.endswith('docPr'):
                                    docPr = d.attrib
                                    break
                            if docPr:
                                print(f"  - docPr: id={docPr.get('id')}, name={docPr.get('name')}, descr={docPr.get('descr')}")
                            
                            # Check srcRect (Cropping)
                            srcRect = None
                            for d in child.iterdescendants():
                                if d.tag.endswith('srcRect'):
                                    srcRect = d.attrib
                                    break
                            if srcRect:
                                print(f"  - srcRect (Cropping Detected): {srcRect}")
                                
                            # Check Fallback
                            # Often wrapping AlternateContent
                            # But here we are iterating descendants. 
                            # If the drawing IS inside AlternateContent/Choice, we verify parent.
                            parent = child.getparent()
                            while parent is not None:
                                if 'AlternateContent' in parent.tag:
                                    print("  - Inside AlternateContent")
                                    break
                                parent = parent.getparent()

    if count == 0:
        print("No drawings found in tables.")

if __name__ == "__main__":
    inspect_table_drawings("/app/test_doc.docx")
