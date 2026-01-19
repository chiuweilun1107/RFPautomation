
import sys
import os
import docx
from docx.oxml.ns import qn
from collections import Counter

# Add logic to find unhandled tags
def audit_gaps(file_path):
    print(f"Auditing Gaps for: {file_path}")
    doc = docx.Document(file_path)
    
    # 1. XML Tag Inventory (What exists in the doc?)
    tag_counts = Counter()
    
    # We'll traverse body elements
    for element in doc.element.body.iter():
        tag = element.tag
        # Simplify tag name (remove namespace for readability)
        if '}' in tag: tag = tag.split('}')[1]
        tag_counts[tag] += 1
        
    print("\n--- XML Element Inventory (Top 20) ---")
    for tag, count in tag_counts.most_common(20):
        print(f"{tag}: {count}")

    # 2. Check for Specific Missing Features
    print("\n--- Feature Gap Analysis ---")
    
    # A. Fields (instrText, fldChar) - Dynamic content like Page Numbers / Date / Form Fields
    has_fields = tag_counts['instrText'] > 0 or tag_counts['fldChar'] > 0
    print(f"- Complex Fields (instrText/fldChar): {'DETECTED' if has_fields else 'None'} (Count: {tag_counts['instrText'] + tag_counts['fldChar']})")
    if has_fields:
        print("  -> Current parser largely IGNORES dynamic fields (e.g. auto-updating dates, page numbers). Text might be static cached version.")

    # B. Inline Drawings (drawing/object) - Inline charts, shapes
    has_drawings = tag_counts['drawing'] > 0
    print(f"- Inline Drawings/Shapes: {'DETECTED' if has_drawings else 'None'} (Count: {tag_counts['drawing']})")
    if has_drawings:
        print("  -> Current parser handles floating images but might miss inline shapes or charts embedded in text.")

    # C. Hyperlinks
    has_links = tag_counts['hyperlink'] > 0
    print(f"- Hyperlinks: {'DETECTED' if has_links else 'None'} (Count: {tag_counts['hyperlink']})")
    
    # D. Comments
    has_comments = tag_counts['commentRangeStart'] > 0
    print(f"- Comments: {'DETECTED' if has_comments else 'None'}")

    # E. Detailed Border Audit
    print("\n--- Table Border Audit ---")
    tables_with_custom_borders = 0
    total_borders_found = 0
    
    for tbl in doc.tables:
        tblPr = tbl._tbl.tblPr
        has_custom = False
        if tblPr is not None:
            tblBorders = tblPr.find(qn('w:tblBorders'))
            if tblBorders is not None:
                has_custom = True
                total_borders_found += 1
        
        # Check cell borders
        for row in tbl.rows:
            for cell in row.cells:
                tcPr = cell._tc.tcPr
                if tcPr is not None and tcPr.find(qn('w:tcBorders')):
                    has_custom = True
                    break
            if has_custom: break
        
        if has_custom:
            tables_with_custom_borders += 1
            
    print(f"- Tables with Custom Borders: {tables_with_custom_borders} / {len(doc.tables)}")
    if tables_with_custom_borders > 0:
        print("  -> Current parser IGNORES custom borders (renders standard grid). Visually distinct borders (bold, double, dashed) will be lost.")

    # F. Run Content Audit (What extract_run_content handles)
    print("\n--- Run Content Analysis ---")
    # We iterate a sample of runs to see if they contain unhandled children
    unhandled_run_children = Counter()
    handled_tags = ['t', 'tab', 'br', 'cr', 'sym', 'rPr'] # rPr is style, handled separately
    
    for p in doc.paragraphs:
        for r in p.runs:
            for child in r._element:
                tag_raw = child.tag
                tag_name = tag_raw.split('}')[1] if '}' in tag_raw else tag_raw
                if tag_name not in handled_tags:
                    unhandled_run_children[tag_name] += 1
                    
    if unhandled_run_children:
        print("Unhandled Run Child Elements:")
        for tag, count in unhandled_run_children.items():
            print(f"  - {tag}: {count}")
    else:
        print("All run content elements are handled or irrelevant.")

if __name__ == "__main__":
    audit_gaps("/app/test_doc.docx")
