
import sys
import os
import json
import docx
from docx.oxml.ns import qn

sys.path.append("/app")

try:
    from service import detect_fillable_tables, extract_images_from_doc
except ImportError as e:
    print(f"Import Error: {e}")
    sys.exit(1)

def verify_images(file_path):
    print(f"Verifying Inline Images for: {file_path}")
    doc = docx.Document(file_path)
    
    # 1. Extract Images (Simulate pipeline)
    print("Extracting images...")
    # NOTE: In verify env, supabase creds might be missing, so this returns empty list usually
    # BUT we need to mock it or hope the logic allows empty/mocked images.
    # Actually, extract_images_from_doc needs credentials.
    # If creds missing, it returns [].
    # If [] returned, then extract_runs will find NOTHING.
    # So we MUST mock the images list or allow extraction.
    
    # Let's inspect doc manually first to see if temp_id are set?
    # No, temp_id are set BY extract_images_from_doc.
    # So we MUST run extract_images_from_doc.
    
    # Set dummy env if needed?
    os.environ["SUPABASE_URL"] = "https://mock.supabase.co"
    os.environ["SUPABASE_SERVICE_KEY"] = "mock_key"
    
    # Mock supabase client creation in service.py? 
    # It imports `create_client`.
    # I can't easily mock imports here without mocking library.
    # But `verify_fidelity.py` ran fine? No, it used `extract_runs(para)` which didn't depend on supabase.
    # This test DOES depend on it.
    
    # Hack: I will just instantiate `images` list manually by inspecting XML and mocking the dict.
    # Or, rely on the fact that `extract_images_from_doc` logs "Supabase init failed" and returns list?
    # No, line 195: returns [] if no creds.
    
    # I will set Mock Env, and hope `create_client` works or fails gracefully?
    # Supabase client usually checks connection only on request.
    
    images = extract_images_from_doc(doc, "test_template")
    # If uploading, it will fail.
    # I need to mock the `upload` method.
    
    # Actually, for verification, I just want to see if `extract_runs` *attempts* to find the image.
    # I can mock the `images` list myself.
    # Inspect the DOCX XML to find a drawing, get its ID.
    # Create an image dict for it.
    # Pass to detect_fillable_tables.
    
    # Find a drawing in a table
    found_drawing_id = None
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                         for child in run._element.iterdescendants():
                             if child.tag.endswith('drawing'):
                                 # Find blip embed
                                 for d in child.iterdescendants():
                                     if d.tag.endswith('blip'):
                                         for k,v in d.attrib.items():
                                             if 'embed' in k:
                                                 found_drawing_id = v
                                                 # Set temp_id manually as extract_images would
                                                 child.set('temp_img_id', 'mock_img_123')
                                                 print(f"Found drawing {v} in table. Mocking...")
    
    if found_drawing_id:
        mock_images = [{
            "id": "mock_img_123",
            "url": "http://mock.url/image.png",
            "width": 100,
            "height": 100,
            "para_temp_id": "mock_para_id" # Not needed for inline?
        }]
        
        # Run detection
        tables = detect_fillable_tables(doc, mock_images)
        
        # Check output
        count = 0
        for tbl in tables:
            for row in tbl['rows']:
                for cell in tbl['cells']:
                    for run in cell['runs']:
                        if run.get('type') == 'image':
                            print(f"SUCCESS: Found Inline Image Run! ID: {run['image_data']['id']}")
                            count += 1
        
        if count == 0:
            print("FAILURE: Drawing found in XML but not extracted as Image Run.")
    else:
        print("No drawings found in tables in this doc.")

if __name__ == "__main__":
    verify_images("/app/test_doc.docx")
