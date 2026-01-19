"""
Word 範本解析服務
功能: 自動識別 Word 文件中的可填寫欄位,生成動態表單 Schema
"""

import os
import re
import json
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, RGBColor, Length
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.shape import WD_INLINE_SHAPE
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml.ns import qn
from docx.text.run import Run
from supabase import create_client, Client

app = FastAPI(title="Template Parsing Service")

# ==================== 數據模型 ====================

class FieldSchema(BaseModel):
    """表單欄位 Schema"""
    name: str
    label: str
    type: str
    required: bool = True
    placeholder: str = ""
    default_value: str = ""
    position: Dict[str, Any] = {}
    style: Dict[str, Any] = {}

class TableFieldSchema(BaseModel):
    """表格欄位 Schema"""
    name: str
    label: str
    columns: List[Dict[str, str]]
    min_rows: int = 1
    max_rows: int = 100
    style: Dict[str, Any] = {}
    position: Dict[str, Any] = {}

class SectionSchema(BaseModel):
    """章節 Schema"""
    index: int
    page_width: float
    page_height: float
    orientation: str
    margin_top: float
    margin_bottom: float
    margin_left: float
    margin_right: float
    header_distance: float
    footer_distance: float

class ImageSchema(BaseModel):
    """圖片 Schema"""
    id: str
    url: str
    width: float
    height: float
    index: int
    paragraph_index: int  # 圖片所在的段落索引
    content_type: str = "image/png"

class StructureItem(BaseModel):
    """結構項目 Schema"""
    type: str  # paragraph, table, image
    id: str    # 對應 field_name, table_name, 或 image_id
    block_index: int # 在 type 中的索引 (例如第幾個段落)

class TemplateSchema(BaseModel):
    """範本 Schema"""
    template_id: str
    template_name: str
    sections: List[SectionSchema]
    fields: List[FieldSchema]
    tables: List[TableFieldSchema]
    images: List[ImageSchema]
    structure: List[Dict[str, Any]] # 完整文件結構順序
    styles: Dict[str, Any]

# ==================== 核心解析邏輯 ====================

def get_length_in_points(length_obj) -> Optional[float]:
    if length_obj is None:
        return None
    if isinstance(length_obj, Length):
        return length_obj.pt
    return float(length_obj) / 12700.0

def get_local_tag(tag):
    return tag.split('}')[-1] if '}' in tag else tag

def extract_section_properties(doc: Document) -> List[Dict[str, Any]]:
    sections_data = []
    for i, section in enumerate(doc.sections):
        width_pt = section.page_width.pt if section.page_width else 0
        height_pt = section.page_height.pt if section.page_height else 0
        orientation = "LANDSCAPE" if width_pt > height_pt else "PORTRAIT"

        sections_data.append({
            "index": i,
            "page_width": width_pt,
            "page_height": height_pt,
            "orientation": orientation,
            "margin_top": section.top_margin.pt if section.top_margin else 0,
            "margin_bottom": section.bottom_margin.pt if section.bottom_margin else 0,
            "margin_left": section.left_margin.pt if section.left_margin else 0,
            "margin_right": section.right_margin.pt if section.right_margin else 0,
            "header_distance": section.header_distance.pt if section.header_distance else 0,
            "footer_distance": section.footer_distance.pt if section.footer_distance else 0,
        })
    return sections_data

def extract_paragraph_style(para) -> Dict[str, Any]:
    style = {}
    if para.runs:
        run = para.runs[0]
        style["font_name"] = run.font.name
        if hasattr(run, 'element') and run.element is not None:
             rPr = run.element.rPr
             if rPr is not None:
                 rFonts = rPr.find(qn('w:rFonts'))
                 if rFonts is not None:
                     east_asia = rFonts.get(qn('w:eastAsia'))
                     if east_asia:
                         style["font_name_cjk"] = east_asia

        style["font_size"] = run.font.size.pt if run.font.size else 12
        style["bold"] = run.font.bold or False
        style["italic"] = run.font.italic or False
        style["underline"] = True if run.font.underline else False
        
        if run.font.color and run.font.color.rgb:
            rgb = run.font.color.rgb
            style["color"] = f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
    
    style["alignment"] = para.alignment.name.lower() if para.alignment else "left"

    pf = para.paragraph_format
    style["indentation"] = {
        "left": pf.left_indent.pt if pf.left_indent else 0,
        "right": pf.right_indent.pt if pf.right_indent else 0,
        "first_line": pf.first_line_indent.pt if pf.first_line_indent else 0,
    }
    
    # Standardize field names for frontend (line -> line_spacing)
    line_spacing = pf.line_spacing if pf.line_spacing else 1.0
    line_spacing_rule = pf.line_spacing_rule.name if pf.line_spacing_rule else "SINGLE"
    
    style["spacing"] = {
        "before": pf.space_before.pt if pf.space_before else 0,
        "after": pf.space_after.pt if pf.space_after else 0,
        "line_spacing": line_spacing,
        "line_spacing_rule": line_spacing_rule,
        "line": line_spacing, # Duplicate for backward compatibility
        "lineRule": line_spacing_rule.lower() # Duplicate for backward compatibility/consistency
    }
    style["pagination"] = {
        "keep_together": pf.keep_together or False,
        "keep_with_next": pf.keep_with_next or False,
        "page_break_before": pf.page_break_before or False,
        "widow_control": pf.widow_control or False
    }

    return style

def check_page_breaks(para) -> bool:
    if para.paragraph_format.page_break_before:
        return True
    for run in para.runs:
        if 'w:br' in run.element.xml and 'type="page"' in run.element.xml:
             return True
        brs = run.element.findall(qn('w:br'))
        for br in brs:
            if br.get(qn('w:type')) == 'page':
                return True
    return False

def extract_images_from_doc(doc: Document, template_id: str) -> List[Dict[str, Any]]:
    images = []
    supabase_url = os.environ.get("SUPABASE_URL")
    supabase_key = os.environ.get("SUPABASE_SERVICE_KEY")
    
    if not supabase_url or not supabase_key:
        print("Warning: Supabase credentials missing. Image extraction skipped.")
        return []

    # Map paragraph objects to their index for quick lookup
    # We will use ID or element reference to match better
    
    try:
        supabase: Client = create_client(supabase_url, supabase_key)
        bucket_name = "generated-documents"
        
        # Helper to get local tag name without namespace
        def get_local_tag(tag):
             return tag.split('}')[-1]
        # Iterate ALL descendants to find potential image containers
        # We look for 'drawing' (modern) and 'pict' (legacy/fallback)
        for i, elem in enumerate(doc.element.body.iterdescendants()):
            tag = get_local_tag(elem.tag)
            
            if tag not in ['drawing', 'pict']:
                continue
                
            try:
                rId = None
                width_pt = 0
                height_pt = 0
                
                if tag == 'drawing':
                    # Find blip (ignoring namespace via tag name)
                    for child in elem.iterdescendants():
                        if get_local_tag(child.tag) == 'blip':
                            # Check attribs aggressively
                            for k, v in child.attrib.items():
                                if k.endswith('embed') or k.endswith('link') or k.endswith('id'):
                                    rId = v
                                    break
                            if rId: break
                            
                    # Dimensions
                    for child in elem.iterdescendants():
                        if get_local_tag(child.tag) == 'extent':
                            cx = int(child.get('cx', 0))
                            cy = int(child.get('cy', 0))
                            width_pt = cx / 12700
                            height_pt = cy / 12700
                            break
                            
                elif tag == 'pict':
                    # Find imagedata
                    for child in elem.iterdescendants():
                        if get_local_tag(child.tag) == 'imagedata':
                             for k, v in child.attrib.items():
                                if k.endswith('id'):
                                    rId = v
                                    break
                             if rId: break
                    
                    # Dimensions logic omitted for brevity, handled by frontend auto-size

                if not rId or rId not in doc.part.related_parts:
                    continue
                    
                image_part = doc.part.related_parts[rId]
                image_bytes = image_part.blob
                content_type = image_part.content_type
                
                ext = content_type.split('/')[-1] if '/' in content_type else 'png'
                filename = f"parsed_image_{uuid.uuid4()}.{ext}"
                path = f"template_assets/{template_id}/{filename}"
                
                supabase.storage.from_(bucket_name).upload(path, image_bytes, {
                    "content-type": content_type,
                    "upsert": "false"
                })
                
                public_url = supabase.storage.from_(bucket_name).get_public_url(path)
                
                # Find parent paragraph using tag check instead of isinstance(CT_P)
                # This avoids import/proxy class issues
                p_element = elem
                while p_element is not None and not p_element.tag.endswith('}p'):
                    p_element = p_element.getparent()
                
                # Mark paragraph with temp_id for robust matching
                para_temp_id = None
                alignment = None
                if p_element is not None:
                    para_temp_id = p_element.get('temp_id')
                    if not para_temp_id:
                        para_temp_id = str(uuid.uuid4())
                        p_element.set('temp_id', para_temp_id)
                    
                    # Try to find alignment from pPr -> jc
                    # p_element is a CT_P
                    pPr = p_element.find(qn('w:pPr'))
                    if pPr is not None:
                         jc = pPr.find(qn('w:jc'))
                         if jc is not None:
                             alignment = jc.get(qn('w:val'))

                # Find matching paragraph index in doc.paragraphs
                para_index = -1
                if p_element is not None:
                    for idx, para in enumerate(doc.paragraphs):
                        # Use temp_id for matching
                        if para._element.get('temp_id') == para_temp_id:
                            para_index = idx
                            # Robust Alignment Extraction using python-docx
                            # This handles direction formatting + style inheritance
                            try:
                                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                                    alignment = 'center'
                                elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                                    alignment = 'right'
                                elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                                    alignment = 'justify'
                                elif para.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                                    alignment = 'left'
                            except:
                                pass
                            break
                
                # Check for Floating Alignment (wp:anchor -> wp:positionH -> wp:align)
                # elem is drawing/pict.
                is_floating = False
                floating_align = None
                try:
                    # Search specifically for wp:anchor
                    # drawing -> anchor
                    for child in elem.iterdescendants():
                         if child.tag.endswith('anchor'):
                             is_floating = True
                             # Look for align
                             for grand in child.iterdescendants():
                                 if grand.tag.endswith('align'):
                                     if grand.text in ['center', 'right', 'left']:
                                         floating_align = grand.text
                                     break
                             break
                except:
                    pass

                # If floating align found, OVERRIDE paragraph alignment
                if floating_align:
                    alignment = floating_align

                img_id = str(uuid.uuid4())
                
                # IMPORTANT: Tag the element with this ID so we can find it during structure traversal
                # elem is 'drawing' or 'pict'
                elem.set('temp_img_id', img_id)
                if is_floating:
                     elem.set('is_floating', 'true')

                images.append({
                    "id": img_id,
                    "url": public_url,
                    "width": width_pt,
                    "height": height_pt,
                    "index": len(images), 
                    "paragraph_index": para_index,
                    "para_temp_id": para_temp_id, 
                    "format": {"alignment": alignment} if alignment else {},
                    "content_type": content_type,
                    "is_floating": is_floating
                })
                print(f"Uploaded image to {public_url} (para_index: {para_index}, temp_id: {para_temp_id})")
                
            except Exception as e:
                print(f"Failed to process image element {tag}: {e}")
                import traceback
                traceback.print_exc()
                continue

                    
    except Exception as e:
        print(f"Supabase init failed: {e}")
        
    return images

def detect_fillable_fields(doc: Document) -> List[Dict[str, Any]]:
    # ... (Keep existing logic, omitted for brevity if unchanged, but included below for safety)
    # Actually user said "parsed_fields" is empty too. Let's keep existing logic but maybe it's too strict?
    # User's template might just be normal text.
    # For now we reuse the existing logic, but structure builder will handle the rest.
    fields = []
    field_counter = 1
    
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        para_style = extract_paragraph_style(para)
        
        has_break = check_page_breaks(para)
        if has_break:
            para_style["has_page_break"] = True
        
        # Original regex logic
        if re.search(r'_{3,}', text):
            match = re.match(r'(.+?)[::]?\s*_{3,}', text)
            label = match.group(1).strip() if match else f"欄位 {field_counter}"
            fields.append({
                "name": f"field_{field_counter}",
                "label": label,
                "type": "text",
                "required": True,
                "placeholder": f"請輸入{label}",
                "position": {"paragraph_index": para_idx, "pattern": "underline", "has_page_break": has_break},
                "style": para_style
            })
            field_counter += 1
        elif re.search(r'【\s*】|〔\s*〕', text):
            match = re.match(r'(.+?)[::]?\s*【\s*】', text)
            label = match.group(1).strip() if match else f"欄位 {field_counter}"
            fields.append({
                "name": f"field_{field_counter}",
                "label": label,
                "type": "text",
                "required": True,
                "placeholder": f"請輸入{label}",
                "position": {"paragraph_index": para_idx, "pattern": "bracket", "has_page_break": has_break},
                "style": para_style
            })
            field_counter += 1
        elif re.search(r'(請填寫|填寫說明|說明[::])', text, re.IGNORECASE):
            label = text.replace("請填寫", "").replace("填寫說明", "").replace(":", "").strip()
            fields.append({
                "name": f"field_{field_counter}",
                "label": label or f"說明 {field_counter}",
                "type": "textarea",
                "required": True,
                "placeholder": f"請輸入{label}",
                "position": {"paragraph_index": para_idx, "pattern": "keyword", "has_page_break": has_break},
                "style": para_style
            })
            field_counter += 1
    
    return fields

def detect_fillable_tables(doc: Document, images: List[Dict] = []) -> List[Dict[str, Any]]:
    tables_list = []
    
    for table_idx, table in enumerate(doc.tables):
        # Allow single row tables too if needed
        if len(table.rows) < 1: 
            continue
        
        table_style = {
            "style_name": table.style.name if table.style else "Normal Table",
            "alignment": table.alignment,
            "autofit": table.autofit,
            "row_heights": [row.height.pt if row.height else None for row in table.rows],
            "col_widths": [] # Could be filled if needed
        }
            
        header_row = table.rows[0]
        columns = []
        
        # Extract Columns (Headers)
        for cell in header_row.cells:
            col_name = cell.text.strip()
            # Simple heuristic for type
            field_type = "text"
            if re.search(r'(數量|金額|價格|單價)', col_name):
                field_type = "number"
            elif re.search(r'(日期|時間)', col_name):
                field_type = "date"
            
            columns.append({
                "name": re.sub(r'[^a-zA-Z0-9_]', '_', col_name.lower()) if col_name else f"col_{len(columns)}",
                "label": col_name or "Column",
                "type": field_type,
                "type": field_type,
                "width": cell.width.pt if cell.width else 0
            })
            
            # Enhanced width extraction for pct
            width_val = 0
            if cell.width:
                width_val = cell.width.pt
            else:
                try:
                    tc = cell._tc
                    tcPr = tc.tcPr
                    if tcPr is not None:
                        tcW = tcPr.find(qn('w:tcW'))
                        if tcW is not None and tcW.get(qn('w:type')) == 'pct':
                             # 5000 = 100%
                             width_val = f"{int(tcW.get(qn('w:w')))/50.0}%"
                except:
                    pass
            
            columns[-1]['width'] = width_val
        
        # Extract Row Data (Content)
        # We start from row 1 if we consider row 0 as header. 
        # But for structure preservation, maybe we want row 0 as well?
        # Usually 'columns' defines the schema, 'rows' defines the default data.
        # If the table is purely for layout, header might just be the first row of data.
        # Let's assume row 0 is header for schema, and we store ALL content including header in 'rows' if we want full fidelity?
        # Or typical: columns = header, rows = slice(1).
        
        
        # Extract Row Data (Content)
        rows_data = []
        cells_data = [] # New list for detailed cell info
        row_formats = [] # New list for row formatting
        
        # 1. Parse Table-Level Borders (Defaults)
        table_borders_def = {}
        tblPr = table._tbl.tblPr
        if tblPr is not None:
            tblBorders = tblPr.find(qn('w:tblBorders'))
            # Scan for all relevant sides including inside borders
            table_borders_def = extract_borders_generic(tblBorders, keys=['top', 'bottom', 'left', 'right', 'insideH', 'insideV'])
            
        total_rows_count = len(table.rows)

        for r_idx, row in enumerate(table.rows):
            # Extract Row Height
            row_height = None
            row_height_rule = None
            if row.height:
                row_height = row.height.pt
                row_height_rule = row.height_rule # EXACT, AUTO, AT_LEAST
            
            row_formats.append({
                "height": {"value": row_height, "rule": str(row_height_rule)} if row_height else None,
                "isHeader": getattr(row, 'is_header', False) # python-docx doesn't strictly track header rows easily
            })
            
            row_dict = {}
            for c_idx, cell in enumerate(row.cells):
                # ... existing logic ...
                # Map content to column name/id for rows_data
                if c_idx < len(columns):
                    col_key = columns[c_idx]['name']
                    row_dict[col_key] = cell.text.strip()
                
                # Extract detailed cell info for cells_data
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # ColSpan (gridSpan)
                gridSpan = tcPr.find(qn('w:gridSpan'))
                colSpan = int(gridSpan.get(qn('w:val'))) if gridSpan is not None else 1
                
                # VMerge
                vMerge = tcPr.find(qn('w:vMerge'))
                vMergeVal = (vMerge.get(qn('w:val')) or "continue") if vMerge is not None else None
                
                # Heuristic: If vMerge is 'restart' but text matches the cell directly above, treat as 'continue'
                # This fixes tables where Word/Generator marks every cell as 'restart' despite identical content
                if vMergeVal == 'restart' and r_idx > 0:
                    current_txt = cell.text.strip()
                    try:
                        prev_cell = table.rows[r_idx-1].cells[c_idx]
                        prev_txt = prev_cell.text.strip()
                        if current_txt == prev_txt:
                            vMergeVal = "continue"
                    except:
                        pass
                
                cell_format = {
                    "colSpan": colSpan,
                    "vMerge": vMergeVal, # start, continue, or None
                }

                # Cell shading/background
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    if fill and fill != "auto":
                        cell_format["backgroundColor"] = fill

                # Vertical Alignment
                vAlign = tcPr.find(qn('w:vAlign'))
                if vAlign is not None:
                    val = vAlign.get(qn('w:val'))
                    # Map DOCX 'center' to CSS 'middle'
                    if val == 'center':
                        cell_format["vAlign"] = "middle"
                    else:
                        cell_format["vAlign"] = val
                
                # Text Direction (Vertical Text)
                textDirection = tcPr.find(qn('w:textDirection'))
                if textDirection is not None:
                     # e.g. tbRl (top to bottom, right to left) -> vertical-rl
                     # btLr (bottom to top, left to right) -> vertical-lr
                     td_val = textDirection.get(qn('w:val'))
                     if td_val in ['tbRl', 'btLr']:
                         cell_format["textDirection"] = "vertical-rl" # Simplified mapping
                         cell_format["textDirection"] = "vertical-rl" # Simplified mapping
                         cell_format["writingMode"] = "vertical-rl"

                         cell_format["writingMode"] = "vertical-rl"

                # Cell Borders - High Fidelity Logic
                # 1. Start with Default Table Borders based on position
                cell_border_style = {}
                num_cols = len(row.cells)
                
                # Top
                if r_idx == 0:
                    if 'borderTop' in table_borders_def: cell_border_style['borderTop'] = table_borders_def['borderTop']
                else:
                    if 'borderInsideH' in table_borders_def: cell_border_style['borderTop'] = table_borders_def['borderInsideH']
                
                # Bottom
                if r_idx == total_rows_count - 1:
                     if 'borderBottom' in table_borders_def: cell_border_style['borderBottom'] = table_borders_def['borderBottom']
                else:
                     if 'borderInsideH' in table_borders_def: cell_border_style['borderBottom'] = table_borders_def['borderInsideH']
                     
                # Left
                if c_idx == 0:
                     if 'borderLeft' in table_borders_def: cell_border_style['borderLeft'] = table_borders_def['borderLeft']
                else:
                     if 'borderInsideV' in table_borders_def: cell_border_style['borderLeft'] = table_borders_def['borderInsideV']

                # Right
                if c_idx == num_cols - 1:
                     if 'borderRight' in table_borders_def: cell_border_style['borderRight'] = table_borders_def['borderRight']
                else:
                     if 'borderInsideV' in table_borders_def: cell_border_style['borderRight'] = table_borders_def['borderInsideV']
                     
                # 2. Extract Specific Cell Overrides (tcBorders)
                cell_overrides = extract_cell_borders(tcPr)
                cell_border_style.update(cell_overrides)

                if cell_border_style:
                    cell_format["borders"] = cell_border_style

                # Runs extraction for the cell
                cell_runs = []
                cell_h_align = None

                for p_idx, p in enumerate(cell.paragraphs):
                    # 1. Determine Horizontal Alignment (First paragraph wins usually)
                    if cell_h_align is None:
                        # Try helper or direct access
                        align_val = p.alignment
                        if align_val == WD_ALIGN_PARAGRAPH.CENTER:
                            cell_h_align = 'center'
                        elif align_val == WD_ALIGN_PARAGRAPH.RIGHT:
                            cell_h_align = 'right'
                        elif align_val == WD_ALIGN_PARAGRAPH.JUSTIFY:
                            cell_h_align = 'both' # Frontend uses 'both' for justify
                        # Check w:jc fallback
                        try:
                            if not cell_h_align and p._element.pPr:
                                jc = p._element.pPr.find(qn('w:jc'))
                                if jc is not None:
                                    val = jc.get(qn('w:val'))
                                    if val == 'center': cell_h_align = 'center'
                                    elif val == 'right': cell_h_align = 'right'
                                    elif val == 'both': cell_h_align = 'both'
                        except: pass
                    
                    # 2. Handle Line Breaks (Between paragraphs)
                    if p_idx > 0:
                        cell_runs.append({
                            "text": "\n",
                            "type": "text",
                            "format": {}
                        })

                    # 3. Runs Extraction (Supports Inline Images)
                    p_runs = extract_runs(p, images)
                    
                    # 4. Check for Numbering (Bullets)
                    try:
                        if p._element.pPr is not None and p._element.pPr.find(qn('w:numPr')) is not None:
                             bullet_run = {
                                 "text": "• ", 
                                 "format": p_runs[0]['format'] if p_runs else {} 
                             }
                             p_runs.insert(0, bullet_run)
                    except:
                        pass
                        
                    cell_runs.extend(p_runs)

                if cell_h_align:
                    cell_format["hAlign"] = cell_h_align

                cells_data.append({
                    "row": r_idx,
                    "col": c_idx,
                    "text": cell.text.strip(),
                    "runs": cell_runs,
                    "format": cell_format
                })
            
            rows_data.append(row_dict)

        # Extract Table Level Properties
        tblPr = table._tbl.tblPr
        table_width = None
        table_indent = None
        
        if tblPr is not None:
             tblW = tblPr.find(qn('w:tblW'))
             if tblW is not None:
                 w_type = tblW.get(qn('w:type'))
                 w_val = int(tblW.get(qn('w:w'), 0))
                 if w_type == 'dxa':
                     table_width = w_val / 20.0
                 elif w_type == 'pct':
                     table_width = f"{w_val / 50.0}%"
                     pass
                     
             tblInd = tblPr.find(qn('w:tblInd'))
             if tblInd is not None:
                 ind_type = tblInd.get(qn('w:type'))
                 ind_val = int(tblInd.get(qn('w:w'), 0))
                 if ind_type == 'dxa':
                     table_indent = ind_val / 20.0
        
        table_style = {
            "width": table_width,
            "marginLeft": table_indent
        }
        
        tables_list.append({
            "name": f"table_{table_idx + 1}",
            "label": f"表格 {table_idx + 1}",
            "columns": columns,
            "columnWidths": [c['width'] for c in columns], # Extract simple array for frontend
            "rows": rows_data,
            "cells": cells_data,
            "rowFormats": row_formats, # Add row formats
            "min_rows": 1,
            "max_rows": 100,
            "position": {"table_index": table_idx},
            "style": table_style
        })
    
    return tables_list

def extract_run_style(run) -> Dict[str, Any]:
    style = {}
    rPr = run.element.rPr
    if rPr is not None:
        # Font
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
             style["font"] = rFonts.get(qn('w:eastAsia')) or rFonts.get(qn('w:ascii'))
        
        # Size
        sz = rPr.find(qn('w:sz'))
        if sz is not None:
            style["size"] = float(sz.get(qn('w:val'))) / 2
            
        # Color
        color = rPr.find(qn('w:color'))
        if color is not None:
            style["color"] = color.get(qn('w:val'))
            
        # Bold
        if rPr.find(qn('w:b')) is not None:
            style["bold"] = True
            
        # Italic
        if rPr.find(qn('w:i')) is not None:
            style["italic"] = True
            
        # Underline
        u = rPr.find(qn('w:u'))
        if u is not None:
            style["underline"] = True

    return style

def extract_borders_generic(borders_element, keys=['top', 'bottom', 'left', 'right']) -> Dict[str, str]:
    """
    Extract borders from a borders element (tcBorders or tblBorders).
    """
    borders = {}
    if borders_element is None:
        return borders
        
    for side in keys:
        b_el = borders_element.find(qn(f'w:{side}'))
        if b_el is not None:
             val = b_el.get(qn('w:val'))
             if val == 'nil':
                 borders[f'border{side.capitalize()}'] = 'none'
             else:
                 # Width (1/8 pt)
                 sz = int(b_el.get(qn('w:sz'), 4))
                 width_pt = sz / 8.0
                 width_px = 1 if width_pt < 1.0 else (2 if width_pt < 2.5 else 3)
                 
                 # Color
                 color = b_el.get(qn('w:color'), '000000')
                 if color == 'auto': color = '000000'
                 
                 # Style
                 style = 'solid'
                 if 'dot' in val: style = 'dotted'
                 elif 'dash' in val: style = 'dashed'
                 elif 'double' in val: style = 'double'
                 elif 'Gap' in val: style = 'double' # Handle thinThickSmallGap etc
                 
                 borders[f'border{side.capitalize()}'] = f'{width_px}px {style} #{color}'
    return borders

def extract_cell_borders(tcPr) -> Dict[str, str]:
    # Wrapper for backward compatibility or ease check
    if tcPr is None: return {}
    tcBorders = tcPr.find(qn('w:tcBorders'))
    return extract_borders_generic(tcBorders)

    return extract_borders_generic(tcBorders)

def extract_run_elements(run, images=[]) -> List[Dict]:
    elements = []
    if run._element is None: return elements

    for child in run._element:
        tag = child.tag
        if tag == qn('w:t'):
            text = child.text
            if text: elements.append({"type": "text", "content": text})
        
        elif tag == qn('w:tab'):
            elements.append({"type": "text", "content": "\t"})
        
        elif tag == qn('w:br') or tag == qn('w:cr'):
            elements.append({"type": "text", "content": "\n"})
            
        elif tag == qn('w:sym'):
            font = child.get(qn('w:font'))
            char = child.get(qn('w:char'))
            sym_text = ""
            if font == 'Wingdings':
                if char == 'F06F': sym_text = "\u2610" # ☐
                elif char == 'F0FE': sym_text = "\u2611" # ☑
                elif char == 'F0A8': sym_text = "\u2611" # ☑
                elif char == 'F078': sym_text = "\u2612" # ☒
            if sym_text:
                 elements.append({"type": "text", "content": sym_text})
        
        elif tag == qn('w:drawing') or child.get('temp_img_id'):
             img_id = child.get('temp_img_id')
             if not img_id:
                 for d in child.iterdescendants():
                     if d.get('temp_img_id'):
                         img_id = d.get('temp_img_id')
                         break
             
             if img_id and images:
                  found = next((img for img in images if img['id'] == img_id), None)
                  if found:
                      elements.append({
                          "type": "image",
                          "image_data": found
                      })

    # Fallback if no elements found but run.text exists
    if not elements and run.text:
         elements.append({"type": "text", "content": run.text})
         
    return elements

def extract_runs(para, images=[]) -> List[Dict[str, Any]]:
    runs = []
    for run in para.runs:
        elements = extract_run_elements(run, images)
        format_style = extract_run_style(run)
        
        current_text = ""
        
        for el in elements:
            if el['type'] == 'text':
                current_text += el['content']
            elif el['type'] == 'image':
                # Flush text
                if current_text:
                    runs.append({"text": current_text, "format": format_style})
                    current_text = ""
                
                # Append Image
                runs.append({
                    "type": "image",
                    "image_data": el['image_data'],
                    "format": format_style
                })
        
        # Flush remaining text
        if current_text:
            runs.append({"text": current_text, "format": format_style})
            
    return runs

def build_document_structure(doc: Document, fields: List[Dict], tables: List[Dict], images: List[Dict]) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    """
    建立文件的線性結構 (包含 Field, Table, Image, AND Plain Paragraphs)
    修復: 確保所有內容都被包含
    Returns:
        structure:List, paragraphs:List
    """
    structure = []
    paragraphs = []
    para_idx = 0
    table_idx = 0
    
    # 遍歷 body 的所有子元素
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            # 這是一個段落
            para_id = str(uuid.uuid4())
            
            # Check for fields (Partial match inside paragraph?)
            related_fields = [f for f in fields if f['position']['paragraph_index'] == para_idx]
            
            # Match images by temp_id
            c_temp_id = child.get('temp_id')
            related_images = []
            if c_temp_id:
                related_images = [img for img in images if img.get('para_temp_id') == c_temp_id]
            else:
                 # Fallback to paragraph_index if temp_id missing
                 related_images = [img for img in images if img['paragraph_index'] == para_idx]

            # If it's a field paragraph, just output the field
            related_fields = [f for f in fields if f['position']['paragraph_index'] == para_idx]
            if related_fields:
                for field in related_fields:
                    structure.append({
                        "type": "field",
                        "id": field['name'],
                        "label": field['label'],
                        "block_type": "field",
                        "block_index": para_idx
                    })
                para_idx += 1
                continue

            # Advanced: Interleave Text and Images based on run order
            # We iterate through the paragraph's XML children (runs)
            current_text_runs = []
            current_uuid = str(uuid.uuid4())
            deferred_images = []
            
            # Helper to commit current text block
            def commit_text_block(runs_list):
                if not runs_list: return
                text_content = "".join([r['text'] for r in runs_list])
                # Only skip if empty and we have images (pure spacer validation handled outside)
                if not text_content.strip() and not related_images:
                    return 
                
                # Simplified: Create a paragraph object
                nonlocal current_uuid
                paragraphs.append({
                    "id": current_uuid,
                    "text": text_content,
                    "style": doc.paragraphs[para_idx].style.name if doc.paragraphs[para_idx].style else "Normal",
                    "format": extract_paragraph_style(doc.paragraphs[para_idx]),
                    "runs": runs_list,
                    "index": para_idx
                })
                structure.append({
                    "type": "paragraph",
                    "id": current_uuid,
                    "block_type": "paragraph",
                    "block_index": para_idx
                })
                current_uuid = str(uuid.uuid4()) # Reset for next block
            
            # Helper to find image by temp_img_id
            def find_image_by_xml_id(xml_id):
                for img in related_images:
                    if img.get('id') == xml_id: return img
                return None

            # Helper to extract runs from a specific XML element
            def extract_runs_from_element(element, parent_para):
                try:
                    run = Run(element, parent_para)
                    text_val = run.text
                    style_val = {}
                    if run.bold: style_val['bold'] = True
                    if run.italic: style_val['italic'] = True
                    if run.underline: style_val['underline'] = True
                    if run.font.color and run.font.color.rgb:
                        style_val['color'] = str(run.font.color.rgb)
                    if run.font.size:
                        style_val['size'] = run.font.size.pt
                    
                    # Capture exact fonts
                    style_val['font'] = run.font.name
                    if hasattr(run, 'element') and run.element is not None:
                        rPr = run.element.rPr
                        if rPr is not None:
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is not None:
                                east_asia = rFonts.get(qn('w:eastAsia'))
                                if east_asia:
                                    style_val['fontCJK'] = east_asia

                    return [{
                        "text": text_val,
                        "format": style_val
                    }]
                except Exception as ex:
                    print(f"Failed to extract run: {ex}")
                    return []


            has_content = False
            try:
                # Iterate direct children of the paragraph (usually w:r or drawings)
                for i, elem in enumerate(child.iterchildren()):
                    # Debug log
                    tag_name = get_local_tag(elem.tag)
                    
                    # Handle Fields
                    if elem.tag == qn('w:fldSimple'):
                        runs = extract_runs_from_element(elem, doc.paragraphs[para_idx])
                        if runs:
                            current_text_runs.extend(runs)
                            has_content = True
                        continue

                    # Handle Runs (w:r)
                    if elem.tag == qn('w:r'):
                        # Iterate inside the run for perfect interleaving of text, drawings, and breaks
                        for j, run_child in enumerate(elem.iterchildren()):
                            tag_local = get_local_tag(run_child.tag)
                            
                            # 1. Check for Image in this run element
                            xml_img_id = run_child.get('temp_img_id')
                            if not xml_img_id:
                                for d in run_child.iterdescendants():
                                    if d.get('temp_img_id'):
                                        xml_img_id = d.get('temp_img_id')
                                        break
                            
                            if xml_img_id:
                                found_img = find_image_by_xml_id(xml_img_id)
                                if found_img:
                                    # print(f"[DEBUG] Found Image {xml_img_id} in Para {para_idx} Run {i} Child {j} ({tag_local})")
                                    commit_text_block(current_text_runs)
                                    current_text_runs = []
                                    structure.append({
                                        "type": "image",
                                        "id": found_img['id'],
                                        "url": found_img['url'],
                                        "width": found_img['width'],
                                        "height": found_img['height'],
                                        "block_type": "image",
                                        "block_index": para_idx,
                                        "format": found_img.get('format', {})
                                    })
                                    has_content = True
                                    continue # Processed as image
                                # else:
                                    # print(f"[DEBUG] Image ID found {xml_img_id} but lookup failed in Para {para_idx}")

                            # 2. Check for Text (w:t)
                            if tag_local == 't':
                                # Use python-docx Run to handle styling correctly
                                run_obj = Run(elem, doc.paragraphs[para_idx])
                                style_val = {}
                                if run_obj.bold: style_val['bold'] = True
                                if run_obj.italic: style_val['italic'] = True
                                if run_obj.underline: style_val['underline'] = True
                                if run_obj.font.color and run_obj.font.color.rgb:
                                    style_val['color'] = str(run_obj.font.color.rgb)
                                if run_obj.font.size:
                                    style_val['size'] = run_obj.font.size.pt
                                style_val['font'] = run_obj.font.name
                                
                                # Capture CJK Font
                                if hasattr(run_obj, 'element') and run_obj.element is not None:
                                    rPr = run_obj.element.rPr
                                    if rPr is not None:
                                        rFonts = rPr.find(qn('w:rFonts'))
                                        if rFonts is not None:
                                            east_asia = rFonts.get(qn('w:eastAsia'))
                                            if east_asia: style_val['fontCJK'] = east_asia
                                
                                t_text = run_child.text
                                if t_text:
                                    current_text_runs.append({
                                        "text": t_text,
                                        "format": style_val
                                    })
                                    has_content = True
                            
                            # 3. Check for Page Breaks
                            elif tag_local == 'lastRenderedPageBreak' or (tag_local == 'br' and run_child.get(qn('w:type')) == 'page'):
                                commit_text_block(current_text_runs)
                                current_text_runs = []
                                structure.append({
                                    "type": "page_break",
                                    "id": str(uuid.uuid4())
                                })
                            
                            # 4. Check for carriage returns / breaks
                            elif tag_local == 'br':
                                current_text_runs.append({"text": "\n", "format": {}})
                            elif tag_local == 'tab':
                                current_text_runs.append({"text": "\t", "format": {}})

                    # Handle direct Drawings
                    elif elem.tag == qn('w:drawing') or elem.get('temp_img_id'):
                        xml_img_id = elem.get('temp_img_id')
                        if not xml_img_id:
                            for d in elem.iterdescendants():
                                if d.get('temp_img_id'):
                                    xml_img_id = d.get('temp_img_id')
                                    break
                        
                        if xml_img_id:
                            found_img = find_image_by_xml_id(xml_img_id)
                            if found_img:
                                print(f"[DEBUG] Found Direct Drawing {xml_img_id} in Para {para_idx}")
                                commit_text_block(current_text_runs)
                                current_text_runs = []
                                structure.append({
                                    "type": "image",
                                    "id": found_img['id'],
                                    "url": found_img['url'],
                                    "width": found_img['width'],
                                    "height": found_img['height'],
                                    "block_type": "image",
                                    "block_index": para_idx,
                                    "format": found_img.get('format', {})
                                })
                                has_content = True

            except Exception as e:
                print(f"Error parsing paragraph interleaved content: {e}")
                # Strong fallback
                if not has_content:
                    runs = extract_runs(doc.paragraphs[para_idx])
                    commit_text_block(runs)
                    has_content = True

            # Commit remaining text
            commit_text_block(current_text_runs)
            
            # If still no content (truly empty paragraph), add a spacer paragraph
            if not has_content:
                paragraphs.append({
                    "id": current_uuid,
                    "text": "",
                    "style": doc.paragraphs[para_idx].style.name if doc.paragraphs[para_idx].style else "Normal",
                    "format": extract_paragraph_style(doc.paragraphs[para_idx]),
                    "runs": [],
                    "index": para_idx
                })
                structure.append({
                    "type": "paragraph",
                    "id": current_uuid,
                    "block_type": "paragraph",
                    "block_index": para_idx
                })

            para_idx += 1
            
        elif isinstance(child, CT_Tbl):
            # 這是一個表格
            # Find matching table object
            # Our tables list is built from doc.tables, which iterates in order.
            related_table = next((t for t in tables if t['position']['table_index'] == table_idx), None)
            
            if related_table:
                # 檢查表格內的分頁符 (lastRenderedPageBreak 或 w:br type="page")
                # 策略: 掃描每一列，找出分頁符所在的列索引
                row_break_indices = []
                # Use raw XML iteration for speed and access to breaks
                for r_idx, row in enumerate(child.iterchildren()):
                    if row.tag.endswith('tr'):
                        # Check XML of the row for page break
                        # Note: This finds break ANYWHERE in the row (cells).
                        # We assume break means "End of Page is reached during or after this row"
                        # So we split separate chunks.
                        # However, iterating ALL content might be slow for huge tables?
                        # But necessary.
                        found_break = False
                        # Optimization: check text runs inside cells? 
                        # Simple XML check:
                        # row_xml = getattr(row, 'xml', None) 
                        
                        # Use iterdescendants to find tags
                        found_break = False
                        for d in row.iterdescendants():
                            if d.tag == qn('w:lastRenderedPageBreak'):
                                found_break = True
                                break
                            if d.tag == qn('w:br') and d.get(qn('w:type')) == 'page':
                                found_break = True
                                break
                        
                        if found_break:
                            row_break_indices.append(r_idx)
                
                if not row_break_indices:
                    # No breaks, keep as is
                    structure.append({
                        "type": "table",
                        "id": related_table['name'],
                        "label": related_table['label'],
                        "columns": related_table['columns'], 
                        "block_type": "table",
                        "block_index": table_idx
                    })
                else:
                    # Split table into chunks
                    # e.g. breaks at [4, 8] -> chunks: 0-4, 5-8, 9-end
                    # Wait, if break is in row 4, does row 4 belong to page 1? Yes.
                    splits = []
                    start = 0
                    for brk_idx in row_break_indices:
                        end = brk_idx + 1 # Include the row with break
                        splits.append((start, end))
                        start = end
                    # Add final chunk
                    total_rows = len(related_table.get('rows', []))
                    if start < total_rows:
                        splits.append((start, total_rows))
                    
                    # Create new table parts
                    for i, (s, e) in enumerate(splits):
                        if s >= e: continue
                        
                        part_name = f"{related_table['name']}_part_{i+1}"
                        part_label = f"{related_table['label']} ({i+1})"
                        
                        # Slice rows
                        part_rows = related_table.get('rows', [])[s:e]
                        
                        # Slice rowFormats
                        part_row_formats = related_table.get('rowFormats', [])[s:e]
                        
                        # Slice and Re-index Cells
                        # original cells have 'row' attribute absolute.
                        # We need to filter and shift 'row' to 0-based for the new table part
                        part_cells = []
                        for cell in related_table.get('cells', []):
                            if s <= cell['row'] < e:
                                new_cell = cell.copy()
                                new_cell['row'] = cell['row'] - s
                                part_cells.append(new_cell)
                        
                        new_table = related_table.copy()
                        new_table['name'] = part_name
                        new_table['label'] = part_label
                        new_table['rows'] = part_rows
                        new_table['cells'] = part_cells
                        new_table['rowFormats'] = part_row_formats # Assign sliced formats
                        new_table['is_part'] = True
                        new_table['part_index'] = i
                        
                        # Register new table
                        tables.append(new_table)
                        
                        # Add to structure
                        structure.append({
                            "type": "table",
                            "id": part_name,
                            "label": part_label,
                            "columns": related_table['columns'], 
                            "block_type": "table",
                            "block_index": table_idx
                        })
                        
                        # Add Page Break if not the last chunk
                        if i < len(splits) - 1:
                            structure.append({
                                "type": "page_break",
                                "id": str(uuid.uuid4())
                            })

            else:
                 # Should not happen if we include all tables above, but safe fallback
                 pass
                 
            table_idx += 1
            
    return structure, paragraphs

# ==================== API 端點 ====================

@app.get("/health")
async def health_check():
    return {"status": "healthy", "service": "template-parsing-v2.2 (pict-support)"}

class ParseRequest(BaseModel):
    file_path: str
    bucket: str = "raw-files"
    template_id: Optional[str] = None

# 共用的解析邏輯
def _process_docx_file(file_path: str, filename: str, template_id: str = None) -> Dict[str, Any]:
    try:
        doc = Document(file_path)
        if not template_id:
            template_id = str(uuid.uuid4())
        
        sections = extract_section_properties(doc)
        images = extract_images_from_doc(doc, template_id) # Extract images first
        fields = detect_fillable_fields(doc)
        
        # Pass images to table detection
        tables = detect_fillable_tables(doc, images) 
        
        # 4. 建立線性結構
        structure, paragraphs = build_document_structure(doc, fields, tables, images)

        # Cleanup internal objects before serialization
        # Cleanup internal objects before serialization
        for img in images:
            img.pop('_parent_elem', None)
        
        # Extract default font size
        doc_default_size = None
        try:
            doc_defaults = doc.styles.element.find(qn('w:docDefaults'))
            if doc_defaults is not None:
                rPrDefault = doc_defaults.find(qn('w:rPrDefault'))
                if rPrDefault is not None:
                    rPr = rPrDefault.find(qn('w:rPr'))
                    if rPr is not None:
                        sz = rPr.find(qn('w:sz'))
                        if sz is not None:
                             doc_default_size = int(sz.get(qn('w:val'))) / 2.0
        except Exception:
            pass

        return {
            "template_id": template_id,
            "template_name": filename,
            "doc_default_size": doc_default_size,
            "sections": sections,
            "fields": fields,
            "tables": tables,
            "images": images,
            "structure": structure, 
            "paragraphs": paragraphs, # New Field
            "styles": {
                "default_font": "微軟正黑體",
                "default_size": 12,
                "note": "Enhanced python-docx extraction (v2) with Images & Structure"
            }
        }
    except Exception as e:
        raise e

@app.post("/parse-template")
async def parse_template(file: UploadFile = File(...)):
    """
    解析 Word 範本 (直接上傳檔案)
    """
    try:
        temp_path = f"/tmp/{file.filename}"
        with open(temp_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        try:
            result = _process_docx_file(temp_path, file.filename)
            return JSONResponse(result)
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/parse-from-supabase")
async def parse_from_supabase(request: ParseRequest):
    """
    從 Supabase Storage 下載並解析 Word 範本
    """
    try:
        supabase_url = os.environ.get("SUPABASE_URL")
        supabase_key = os.environ.get("SUPABASE_SERVICE_KEY")
        
        if not supabase_url or not supabase_key:
            raise HTTPException(status_code=500, detail="Supabase credentials missing")

        supabase: Client = create_client(supabase_url, supabase_key)
        
        # 下載檔案
        file_path = request.file_path
        bucket = request.bucket
        
        print(f"Downloading {file_path} from bucket {bucket}...")
        
        try:
            response = supabase.storage.from_(bucket).download(file_path)
            # response is bytes in newer supabase-py versions, or handle checks
        except Exception as e:
            print(f"Failed to download from Supabase: {e}")
            raise HTTPException(status_code=400, detail=f"Failed to download file: {str(e)}")

        # 寫入暫存檔
        filename = os.path.basename(file_path)
        temp_path = f"/tmp/{uuid.uuid4()}_{filename}"
        
        with open(temp_path, "wb") as f:
            f.write(response)
            
        try:
            # 使用傳入的 template_id (若有)，否則 _process_docx_file 會生成
            result = _process_docx_file(temp_path, filename, request.template_id)
            return JSONResponse(result)
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    except HTTPException as he:
        raise he
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8004)
