
import sys
import docx
from docx.oxml.ns import qn

sys.path.append("/app")

def verify_indent(file_path):
    print(f"Verifying Indentation in Tables for: {file_path}")
    doc = docx.Document(file_path)
    
    for t_idx, tbl in enumerate(doc.tables):
        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, p in enumerate(cell.paragraphs):
                    pPr = p._element.pPr
                    if pPr is not None:
                        ind = pPr.find(qn('w:ind'))
                        numPr = pPr.find(qn('w:numPr'))
                        
                        if ind is not None or numPr is not None:
                            text_snippet = p.text[:20]
                            print(f"\n[Table {t_idx+1} R{r_idx}C{c_idx} P{p_idx}] Text: '{text_snippet}...'")
                            
                            if ind is not None:
                                left = ind.get(qn('w:left'))
                                hanging = ind.get(qn('w:hanging'))
                                firstLine = ind.get(qn('w:firstLine'))
                                print(f"  - Indent: left={left}, hanging={hanging}, firstLine={firstLine}")
                                
                            if numPr is not None:
                                ilvl = numPr.find(qn('w:ilvl'))
                                numId = numPr.find(qn('w:numId'))
                                lvl_val = ilvl.get(qn('w:val')) if ilvl is not None else '?'
                                num_val = numId.get(qn('w:val')) if numId is not None else '?'
                                print(f"  - List: ilvl={lvl_val}, numId={num_val}")

    print("\nScan complete.")

if __name__ == "__main__":
    verify_indent("/app/test_doc.docx")
