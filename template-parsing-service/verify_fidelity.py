
import sys
import os
import json
import docx

# Inside Docker, /app is the workdir
sys.path.append("/app")

try:
    from service import extract_runs, detect_fillable_tables
except ImportError as e:
    print(f"Import Error: {e}")
    sys.exit(1)

def verify_fidelity(file_path):
    print(f"Verifying Fidelity for: {file_path}")
    doc = docx.Document(file_path)
    
    print("\n--- Symbol Extraction Verification ---")
    # Find Table 2, Row 9, Col 2 (Index 9, 2) which contains Checkboxes
    # We look for table containing "頁數限制"
    target_table = None
    for tbl in doc.tables:
        # Check first column text roughly
        if len(tbl.rows) > 9 and "頁數限制" in tbl.cell(9, 1).text:
            target_table = tbl
            break
            
    if not target_table:
        print("Could not find table with '頁數限制'. Checking all tables...")
        for idx, tbl in enumerate(doc.tables):
             if len(tbl.rows) > 0:
                 print(f"Table {idx}: {tbl.rows[0].cells[0].text[:10]}...")
        return

    # Check Cell [9, 2]
    cell = target_table.cell(9, 2)
    print(f"Cell Text (Standard): {cell.text}")
    
    extracted_runs = []
    for p in cell.paragraphs:
        extracted_runs.extend(extract_runs(p))
    
    full_text = "".join([r['text'] for r in extracted_runs])
    print(f"Service Extracted Text: {full_text}")
    print(f"Hex Dump: {[hex(ord(c)) for c in full_text]}")
    
    # Check for Unicode Checkboxes
    # \u2610 (Empty Box), \u2611 (Checked Box), \u2612 (X Box)
    if "\u2610" in full_text or "\u2611" in full_text:
        print("SUCCESS: Checkbox detected in final output!")
    else:
        print("FAILURE: Checkbox NOT detected.")

if __name__ == "__main__":
    verify_fidelity("/app/test_doc.docx")
